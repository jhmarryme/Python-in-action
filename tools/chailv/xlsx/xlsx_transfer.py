import pandas as pd
from docx import Document
from docx.shared import Pt
import os

def process_excel_to_word(input_file, output_file):
    # 检查文件是否存在
    if not os.path.exists(input_file):
        print(f"错误：文件 '{input_file}' 不存在。")
        return

    try:
        # 读取Excel文件的两个sheet
        sheet1_df = pd.read_excel(input_file, sheet_name='Sheet1')
        sheet2_df = pd.read_excel(input_file, sheet_name='Sheet2')
    except Exception as e:
        print(f"读取Excel文件出错: {e}")
        return

    # 创建Word文档
    doc = Document()

    # 定义标题格式
    def add_title(doc, text):
        title = doc.add_paragraph()
        run = title.add_run(text)
        run.bold = True
        run.font.size = Pt(12)

    # 处理Sheet1
    doc.add_paragraph("### Sheet1 内容 ###", style='Heading 1')
    for index, row in sheet1_df.iterrows():
        try:
            add_title(doc, f"## 序号: {row['序号']}")
            add_title(doc, f"## 错误码: {row['错误码']}")
            add_title(doc, f"## 名称: {row['名称']}")
            add_title(doc, f"## 错误描述（英文）: {row['错误描述（英文）']}")
            add_title(doc, f"## 错误描述（中文）: {row['错误描述（中文）']}")
            doc.add_paragraph("\n**++**")  # 添加分隔符
        except KeyError as e:
            print(f"警告: Sheet1缺少列 {e}")
        except Exception as e:
            print(f"处理Sheet1第{index+1}行数据出错: {e}")

    # 处理Sheet2
    doc.add_paragraph("### Sheet2 内容 ###", style='Heading 1')
    for index, row in sheet2_df.iterrows():
        try:
            add_title(doc, f"## 序号: {row['序号']}")
            add_title(doc, f"## 接口: {row['接口']}")
            add_title(doc, f"## 错误码: {row['错误码']}")
            add_title(doc, f"## 名称: {row['名称']}")
            add_title(doc, f"## 错误描述（英文）: {row['错误描述（英文）']}")
            add_title(doc, f"## 错误描述（中文）: {row['错误描述（中文）']}")
            doc.add_paragraph("\n**++**")  # 添加分隔符
        except KeyError as e:
            print(f"警告: Sheet2缺少列 {e}")
        except Exception as e:
            print(f"处理Sheet2第{index+1}行数据出错: {e}")

    # 保存Word文档
    try:
        doc.save(output_file)
        print(f"文件成功保存为: {output_file}")
    except Exception as e:
        print(f"保存Word文档出错: {e}")

# 使用示例
process_excel_to_word('test.xlsx', 'output.docx')
