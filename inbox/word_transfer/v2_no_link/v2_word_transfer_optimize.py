import io

from docx import Document
from docx.shared import Inches


def process_word_table_with_inline_images(input_path, output_path_base):
    # 读取原始文档
    doc = Document(input_path)

    # 初始化计数器
    problem_counter = 1
    table_count = 0
    rows_processed = 0
    new_doc = Document()

    for table in doc.tables:
        for row in table.rows[1:]:  # 跳过表头行
            serial_num = row.cells[0].text.strip()
            category = row.cells[1].text.strip()
            user_question = row.cells[2]
            solution = row.cells[3]
            db_operation = row.cells[4]
            print(f'开始处理 {serial_num} 文本: {user_question.text.strip()}')

            # 生成段落块文本并添加问题编号
            block_text = f"**++**\n# 问题 {problem_counter}: \n## 类别：{category}\n ## 用户问题："
            new_doc.add_paragraph(block_text)

            # 处理"user_question"内容并保留图片
            process_paragraph_with_images(user_question, new_doc)

            # 处理"解决思路"内容并保留图片
            if solution.text.strip():
                new_doc.add_paragraph("\n## 解决思路：")
                process_paragraph_with_images(solution, new_doc)
            else:
                new_doc.add_paragraph("无")

            # 处理"数据库操作"内容并保留图片
            if db_operation.text.strip():
                new_doc.add_paragraph("\n## 数据库操作：")
                process_paragraph_with_images(db_operation, new_doc)

            # 控制每100行生成一个新文件
            rows_processed += 1
            if rows_processed >= 100:
                table_count += 1
                new_doc.save(f"{output_path_base}_part{table_count}.docx")
                print(f"处理完成，生成的文档保存在：{output_path_base}_part{table_count}.docx")
                new_doc = Document()  # 重置新的文档
                rows_processed = 0  # 重置行计数器

            # 增加问题计数器
            problem_counter += 1

    # 处理剩余的行
    if rows_processed > 0:
        table_count += 1
        new_doc.save(f"{output_path_base}_part{table_count}.docx")
        print(f"处理完成，生成的文档保存在：{output_path_base}_part{table_count}.docx")


def process_paragraph_with_images(cell, new_doc):
    # 遍历单元格中的段落
    if cell.text.strip():
        for para in cell.paragraphs:
            paragraph = new_doc.add_paragraph()
            for run in para.runs:
                images = run.element.xpath('.//pic:pic')
                if images:
                    for image in images:
                        ids = image.xpath('.//a:blip/@r:embed')
                        for rel in cell.part.rels.values():
                            if "image" in rel.target_ref:
                                if rel.rId in ids:
                                    print(rel.rId)
                                    image_stream = io.BytesIO(rel.target_part.blob)
                                    paragraph.add_run().add_picture(image_stream, width=Inches(2))
                else:
                    paragraph.add_run(run.text)
    else:
        new_doc.add_paragraph("无")


# 指定文件路径
input_path = "test.docx"  # 输入的文件路径
output_path_base = "processed_output_chailv"  # 输出文件路径基准名

# 运行处理函数
process_word_table_with_inline_images(input_path, output_path_base)
