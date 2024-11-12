import io

from docx import Document
from docx.shared import Inches


def process_word_table_with_inline_images(config):
    # 指定文件路径
    input_path = config["input_path"]  # 输入的文件路径
    output_path_base = config["output_path_base"]  # 输出文件路径基准名
    # 读取原始文档
    doc = Document(input_path)

    # 初始化计数器
    problem_counter = 1
    table_count = 0
    rows_processed = 0
    new_doc = Document()

    # 获取配置中的参数
    column_mapping = config["column_mapping"]  # 列映射：列序号 -> 列名
    process_image_columns = config["image_columns"]  # 需要处理图片的列
    rows_per_file = config.get("rows_per_file", 100)  # 每个文件的行数限制
    default_placeholder = config.get("default_placeholder", "无")  # 默认占位文本

    for table in doc.tables:
        for row in table.rows[1:]:  # 跳过表头行
            # 动态提取各列内容
            data = {}
            for col_index, col_name in column_mapping.items():
                data[col_index] = row.cells[col_index]

            # 生成段落块文本并添加问题编号
            block_text = f"\n**++**\n# 序号 {problem_counter}:"
            new_doc.add_paragraph(block_text)

            # 遍历 data 字典的每一列，根据列序号从 column_mapping 获取列名
            for col_index, cell in data.items():
                col_name = column_mapping.get(col_index)  # 从配置的 column_mapping 获取列名

                if col_name:  # 如果列名存在
                    # 添加段落文本，列名通过 col_name 获取
                    new_doc.add_paragraph(f"## {col_name}：")

                    # 判断当前列是否需要处理图片，或者是否有内容
                    text = cell.text.strip()
                    if col_index in process_image_columns:
                        # 如果当前列需要处理图片
                        process_paragraph_with_images(cell, new_doc, process_image_columns)
                    else:
                        # 处理普通文本内容
                        if text:
                            new_doc.add_paragraph(text)
                        else:
                            new_doc.add_paragraph(default_placeholder)

            # 控制每100行生成一个新文件
            rows_processed += 1
            if rows_processed >= rows_per_file:
                table_count += 1
                new_doc.save(f"{output_path_base}_part{table_count}.docx")
                print(f"处理完成，生成的文档保存在：{output_path_base}_part{table_count}.docx")
                new_doc = Document()  # 重置新的文档
                rows_processed = 0  # 重置行计数器

            # 增加问题计数器
            problem_counter += 1

    # 处理剩余的行
    if rows_processed > 0:
        table_count += 1
        new_doc.save(f"{output_path_base}_part{table_count}.docx")
        print(f"处理完成，生成的文档保存在：{output_path_base}_part{table_count}.docx")


def process_paragraph_with_images(cell, new_doc, process_image_columns):
    # 遍历单元格中的段落
    if cell.text.strip():
        for para in cell.paragraphs:
            paragraph = new_doc.add_paragraph()
            for run in para.runs:
                images = run.element.xpath('.//pic:pic')
                if images:
                    for image in images:
                        ids = image.xpath('.//a:blip/@r:embed')
                        for rel in cell.part.rels.values():
                            if "image" in rel.target_ref:
                                if rel.rId in ids:
                                    print(rel.rId)
                                    image_stream = io.BytesIO(rel.target_part.blob)
                                    paragraph.add_run().add_picture(image_stream, width=Inches(2))
                else:
                    paragraph.add_run(run.text)
    else:
        new_doc.add_paragraph("无")


# 配置项
config = {
    "column_mapping": {  # 列名到解析名称的映射
        # 0: "序号",  # 第一列
        1: "类别",  # 第二列
        2: "用户问题",  # 第三列
        3: "解决思路",  # 第四列
        4: "数据库操作"  # 第五列
    },
    "image_columns": [2, 3, 4],  # 包含图片的列的序号
    "rows_per_file": 30,  # 每个文件包含的最大行数
    "input_path": "chailv.docx",  # 输入的文件路径
    "output_path_base": "processed_output_chailv",  # 输出文件路径基准名
    "default_placeholder": "无"  # 默认占位文本
}

# 运行处理函数
process_word_table_with_inline_images(config)
