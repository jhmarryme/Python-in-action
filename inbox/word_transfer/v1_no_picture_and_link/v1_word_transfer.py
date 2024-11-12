from docx import Document

def process_word_table(input_path, output_path):
    # 打开原始文档
    doc = Document(input_path)

    # 创建新的 Word 文档
    new_doc = Document()

    # 遍历文档中的表格
    for table in doc.tables:
        # 遍历表格中的每一行（跳过表头）
        for row in table.rows[1:]:
            cells = row.cells
            # 提取各列数据
            serial_no = cells[0].text.strip()
            category = cells[1].text.strip()
            user_question = cells[2].text.strip()
            solution = cells[3].text.strip()
            db_operation = cells[4].text.strip()

            # 添加格式化段落到新的文档
            new_doc.add_paragraph("**++**")
            new_doc.add_paragraph(f"类别：{category}，用户问题：{user_question}")
            new_doc.add_paragraph("解决思路：")
            new_doc.add_paragraph(solution if solution else "无")
            if db_operation:
                new_doc.add_paragraph("数据库操作：")
                new_doc.add_paragraph(db_operation)
            new_doc.add_paragraph("\n")  # 添加一个空行分隔不同的记录

    # 保存新文档
    new_doc.save(output_path)

# 文件路径（上传的文件路径和输出路径）
input_file_path = "test.docx"
output_file_path = "processed_output.docx"

# 执行函数
process_word_table(input_file_path, output_file_path)
